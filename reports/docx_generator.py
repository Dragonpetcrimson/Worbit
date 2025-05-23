"""
reports/docx_generator.py - DOCX report generation with proper implementation
"""

import os
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import re

# Import path utilities
from utils.path_utils import (
    get_output_path,
    OutputType,
    normalize_test_id,
    get_standardized_filename
)

# Import the python-docx library
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    logging.warning("python-docx library not available. DOCX generation will be limited.")
    DOCX_AVAILABLE = False

from reports.base import ReportGenerator, ReportConfig, ReportData

def extract_root_cause_from_summary(summary: str) -> str:
    """Extract root cause section from the AI-generated summary."""
    root_cause_match = re.search(r'(?:ROOT CAUSE|1\.\s*ROOT CAUSE):\s*([^\n]+(?:\n[^\n\d]+)*)', summary, re.IGNORECASE)
    if root_cause_match:
        return root_cause_match.group(1).strip()
    return "Root cause not identified"

def extract_impact_from_summary(summary: str) -> str:
    """Extract impact section from the AI-generated summary."""
    impact_match = re.search(r'(?:IMPACT|2\.\s*IMPACT):\s*([^\n]+(?:\n[^\n\d]+)*)', summary, re.IGNORECASE)
    if impact_match:
        return impact_match.group(1).strip()
    return "Impact not identified"

def extract_recommended_actions(summary: str) -> List[str]:
    """Extract recommended actions from the AI-generated summary."""
    actions_match = re.search(r'(?:RECOMMENDED ACTIONS|3\.\s*RECOMMENDED ACTIONS):\s*(.*?)(?:\n\n|\Z)', summary, re.DOTALL | re.IGNORECASE)
    if actions_match:
        actions_text = actions_match.group(1).strip()
        # Try to split by bullet points or numbered list
        if '-' in actions_text:
            return [action.strip()[2:].strip() if action.strip().startswith('- ') else action.strip() 
                   for action in actions_text.split('\n-')]
        else:
            return [action.strip() for action in actions_text.split('\n')]
    return ["No specific actions recommended"]

def get_representative_errors(clusters: Dict[int, List[Dict]], max_errors: int = 3) -> List[str]:
    """Extract representative errors from clusters, prioritizing high severity errors."""
    representative_errors = []
    
    # Sort clusters by severity
    sorted_clusters = sorted(
        clusters.items(),
        key=lambda x: sum(1 for err in x[1] if err.get('severity', 'Low') == 'High'),
        reverse=True
    )
    
    for _, errors in sorted_clusters:
        # Sort errors within cluster by severity
        sorted_errors = sorted(
            errors,
            key=lambda x: {'High': 0, 'Medium': 1, 'Low': 2}.get(x.get('severity', 'Low'), 3)
        )
        
        for error in sorted_errors[:2]:  # Take at most 2 errors from each cluster
            if len(representative_errors) < max_errors:
                error_text = error.get('text', '')
                if error_text:
                    representative_errors.append(error_text)
            else:
                break
                
        if len(representative_errors) >= max_errors:
            break
            
    return representative_errors

def generate_bug_document(
    output_dir: str,
    test_id: str,
    summary: str,
    errors: List[Dict],
    ocr_data: List[Dict],
    clusters: Dict[int, List[Dict]],
    background_text: str = "",
    scenario_text: str = "",
    component_analysis: Optional[Dict[str, Any]] = None,
    primary_issue_component: str = "unknown",
    component_report_path: Optional[str] = None
) -> str:
    """
    Generate a DOCX file formatted for Jira bug submission.
    
    Args:
        output_dir: Directory to save the document
        test_id: Test ID for the document title
        summary: Analysis summary text
        errors: List of errors
        ocr_data: OCR text from screenshots
        clusters: Dictionary mapping cluster IDs to lists of errors
        background_text: Background section from feature file
        scenario_text: Scenario section from feature file
        component_analysis: Optional component relationship analysis results
        primary_issue_component: Primary component identified as causing issues
        component_report_path: Path to component report file if available
        
    Returns:
        Path to the generated document
    """
    # Use path utilities
    test_id = normalize_test_id(test_id)
    filename = get_standardized_filename(test_id, "bug_report", "docx")
    doc_path = get_output_path(
        output_dir,
        test_id,
        filename,
        OutputType.PRIMARY_REPORT
    )
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    
    if not DOCX_AVAILABLE:
        # Create a simple text file as a fallback
        with open(doc_path, 'w', encoding='utf-8') as f:
            f.write(f"Bug Report for {test_id}\n\n")
            f.write(f"Summary: {summary}\n\n")
            f.write("Note: Full DOCX generation requires the python-docx library.")
        logging.warning("Using fallback text file for DOCX report due to missing python-docx library")
        return doc_path
    
    try:
        # Create document
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        title_style = styles['Heading 1']
        title_style.font.bold = True
        title_style.font.size = Pt(14)
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add title
        doc.add_heading(f"Bug Report for {test_id}", 0)
        
        # Extract information from summary
        root_cause = extract_root_cause_from_summary(summary)
        impact = extract_impact_from_summary(summary)
        
        # Get representative errors
        rep_errors = get_representative_errors(clusters)
        
        # High severity errors count
        high_sev_count = sum(1 for group in clusters.values() for err in group if err.get('severity', 'Low') == 'High')
        
        # Section 1: Affected Tests
        doc.add_heading("Affected Tests:", 1)
        p = doc.add_paragraph()
        p.add_run(f"{test_id}")
        if high_sev_count > 0:
            p.add_run(f" ({high_sev_count} high severity errors found)")
        
        # Section 2: Details
        doc.add_heading("Details:", 1)
        p = doc.add_paragraph()
        p.add_run(f"{root_cause}\n\n")
        p.add_run(f"Impact: {impact}")
        
        # Add component analysis if available
        if component_analysis:
            try:
                doc.add_heading("Component Analysis:", 1)
                
                # Add root cause component if available
                root_cause_component = primary_issue_component or component_analysis.get("root_cause_component", "unknown")
                if root_cause_component and root_cause_component != "unknown":
                    p = doc.add_paragraph()
                    p.add_run("Root Cause Component: ").bold = True
                    p.add_run(root_cause_component.upper())
                    
                    # Add more component details if available
                    if "component_summary" in component_analysis:
                        for component in component_analysis.get("component_summary", []):
                            if component.get("id") == root_cause_component:
                                if component.get("description"):
                                    p.add_run(f"\nDescription: {component.get('description')}")
                                break
                
                # Add affected components
                affected_components = component_analysis.get("components_with_issues", [])
                if not affected_components and "component_summary" in component_analysis:
                    affected_components = [c.get("id") for c in component_analysis.get("component_summary", [])]
                    
                if affected_components:
                    p = doc.add_paragraph("Affected Components:")
                    for component in affected_components[:5]:  # Limit to first 5
                        if component != "unknown" and component != root_cause_component:
                            p2 = doc.add_paragraph(component.upper(), style='List Bullet')
                
                # Add propagation paths if available
                propagation_paths = component_analysis.get("propagation_paths", [])
                if propagation_paths:
                    p = doc.add_paragraph()
                    p.add_run("Error Propagation Path: ").bold = True
                    for path in propagation_paths[:1]:  # Just show the first path
                        p.add_run(" → ".join([comp.upper() for comp in path]))
            except Exception as e:
                logging.warning(f"Error adding component analysis to document: {str(e)}")

        # Add link to component report if provided
        if component_report_path:
            try:
                report_name = os.path.basename(component_report_path)
                p = doc.add_paragraph()
                p.add_run("Component Analysis Report: ").bold = True
                p.add_run(report_name)
            except Exception as e:
                logging.warning(f"Error adding component report link: {str(e)}")
        
        # Section 3: Log Snippet
        doc.add_heading("Log Snippet:", 1)
        for error in rep_errors:
            doc.add_paragraph(error)
        
        # Section 4: Expected Behavior
        doc.add_heading("Expected Behavior:", 1)
        doc.add_paragraph("The test should complete successfully with all assertions passing.")
        
        # Section 5: Actual Behavior
        doc.add_heading("Actual Behavior:", 1)
        p = doc.add_paragraph()
        p.add_run("Test failed with errors. ")
        p.add_run("See log snippets and details sections for more information.")
        
        # Section 6: Scenario
        doc.add_heading("Scenario:", 1)
        if scenario_text:
            doc.add_paragraph(scenario_text)
        else:
            doc.add_paragraph("No specific scenario information available.")
        
        # Add recommended actions
        actions = extract_recommended_actions(summary)
        if actions and actions[0] != "No specific actions recommended":
            doc.add_heading("Recommended Actions:", 1)
            for action in actions:
                doc.add_paragraph(action, style='List Bullet')
        
        # Add OCR data if available
        if ocr_data:
            doc.add_heading("Additional Information (OCR from Screenshots):", 1)
            for i, ocr in enumerate(ocr_data[:2]):  # Limit to 2 screenshots
                doc.add_paragraph(f"Screenshot {i+1}: {ocr.get('text', '')[:150]}...")
        
        # Save document
        doc.save(doc_path)
        
        logging.info(f"Generated bug report document: {doc_path}")
        return doc_path
        
    except Exception as e:
        logging.error(f"Error generating bug report document: {str(e)}")
        # Create a simple fallback file
        with open(doc_path, 'w', encoding='utf-8') as f:
            f.write(f"Bug Report for {test_id}\n\n")
            f.write(f"Error generating report: {str(e)}\n\n")
            f.write(f"Summary: {summary}\n\n")
        return doc_path

class DocxReportGenerator(ReportGenerator):
    """Generator for DOCX reports."""
    
    def generate(self, data: ReportData) -> str:
        """
        Generate a DOCX report.
        
        Args:
            data: Report data
            
        Returns:
            Path to the generated report
        """
        try:
            # Use path utilities for standardization
            test_id = normalize_test_id(self.config.test_id)
                
            # Generate the document using the function above
            docx_path = generate_bug_document(
                output_dir=self.config.output_dir,
                test_id=test_id,
                summary=data.summary,
                errors=data.errors,
                ocr_data=data.ocr_data,
                clusters=data.clusters,
                background_text=data.background_text,
                scenario_text=data.scenario_text,
                component_analysis=data.component_analysis,
                primary_issue_component=self.config.primary_issue_component,
                component_report_path=(data.component_analysis.get("report_path")
                                      if isinstance(data.component_analysis, dict) else None)
            )
            
            logging.info(f"Bug report document written to: {docx_path}")
            return docx_path
        except Exception as e:
            logging.error(f"Error generating bug report document: {str(e)}")
            return ""